#!/usr/bin/env python3
"""
Word Document Formatter Utility

This module provides methods to convert text content into nicely formatted Word documents
with proper styling, headers, tables, and formatting suitable for ML research reports.
"""

import os
import re
from datetime import datetime
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("⚠️ python-docx not installed. Install with: pip install python-docx")
    exit(1)


@dataclass
class DocumentStyle:
    """Configuration for document styling."""
    title_size: int = 16
    heading1_size: int = 14
    heading2_size: int = 12
    heading3_size: int = 11
    body_size: int = 10
    code_size: int = 9
    font_name: str = "Calibri"
    code_font: str = "Consolas"
    line_spacing: float = 1.15
    margins: float = 1.0  # inches


class WordFormatter:
    """Main class for formatting text content into Word documents."""
    
    def __init__(self, style: Optional[DocumentStyle] = None):
        """Initialize the formatter with styling options."""
        self.style = style or DocumentStyle()
        self.doc = Document()
        self._setup_document_styles()
    
    def _setup_document_styles(self):
        """Set up custom styles for the document."""
        styles = self.doc.styles
        
        # Configure Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.style.font_name
        normal_font.size = Pt(self.style.body_size)
        normal_para = normal_style.paragraph_format
        normal_para.line_spacing = self.style.line_spacing
        normal_para.space_after = Pt(6)
        
        # Create custom heading styles if they don't exist
        self._create_custom_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH, 
                                 self.style.title_size, bold=True, color=RGBColor(0, 51, 102))
        
        self._create_custom_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH, 
                                 self.style.heading1_size, bold=True, color=RGBColor(0, 102, 204))
        
        self._create_custom_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH, 
                                 self.style.heading2_size, bold=True, color=RGBColor(51, 51, 51))
        
        self._create_custom_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH, 
                                 self.style.heading3_size, bold=True, color=RGBColor(102, 102, 102))
        
        self._create_custom_style('Code Block', WD_STYLE_TYPE.PARAGRAPH, 
                                 self.style.code_size, font_name=self.style.code_font)
        
        self._create_custom_style('Highlight', WD_STYLE_TYPE.CHARACTER, 
                                 self.style.body_size, color=RGBColor(0, 102, 0), bold=True)
    
    def _create_custom_style(self, name: str, style_type: int, size: int, 
                           bold: bool = False, color: Optional[RGBColor] = None, 
                           font_name: Optional[str] = None):
        """Create a custom style."""
        try:
            style = self.doc.styles.add_style(name, style_type)
        except ValueError:
            # Style already exists
            style = self.doc.styles[name]
        
        if style_type == WD_STYLE_TYPE.PARAGRAPH:
            font = style.font
            para_format = style.paragraph_format
            para_format.space_after = Pt(6)
            if name == 'Code Block':
                para_format.left_indent = Inches(0.25)
                # Add light gray background for code blocks
                shading_elm = parse_xml(r'<w:shd {} w:fill="F8F8F8"/>'.format(nsdecls('w')))
                style._element.get_or_add_pPr().append(shading_elm)
        else:
            font = style.font
        
        font.name = font_name or self.style.font_name
        font.size = Pt(size)
        font.bold = bold
        if color:
            font.color.rgb = color
    
    def add_title(self, title: str, subtitle: str = "", add_date: bool = True) -> None:
        """Add a formatted title section."""
        # Main title
        title_para = self.doc.add_paragraph(title, style='Custom Title')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle if provided
        if subtitle:
            subtitle_para = self.doc.add_paragraph(subtitle, style='Custom Heading 2')
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        if add_date:
            date_para = self.doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font = date_para.runs[0].font
            font.size = Pt(9)
            font.color.rgb = RGBColor(102, 102, 102)
        
        # Add page break
        self.doc.add_page_break()
    
    def add_table_of_contents(self, headings: List[str]) -> None:
        """Add a simple table of contents."""
        toc_para = self.doc.add_paragraph("Table of Contents", style='Custom Heading 1')
        
        for i, heading in enumerate(headings, 1):
            toc_item = self.doc.add_paragraph(f"{i}. {heading}")
            toc_item.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_page_break()
    
    def add_heading(self, text: str, level: int = 1) -> None:
        """Add a heading with the specified level (1-3)."""
        if level == 1:
            style = 'Custom Heading 1'
        elif level == 2:
            style = 'Custom Heading 2'
        elif level == 3:
            style = 'Custom Heading 3'
        else:
            style = 'Custom Heading 3'
        
        self.doc.add_paragraph(text, style=style)
    
    def add_paragraph(self, text: str, style: str = 'Normal') -> None:
        """Add a regular paragraph."""
        self.doc.add_paragraph(text, style=style)
    
    def add_formatted_paragraph(self, text: str, style: str = 'Normal') -> None:
        """Add a paragraph with markdown-style formatting (bold, italic)."""
        para = self.doc.add_paragraph(style=style)
        
        # Parse and format the text
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = para.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = para.add_run(part[1:-1])
                run.font.italic = True
            else:
                # Normal text
                para.add_run(part)
    
    def add_bullet_list(self, items: List[str], style: str = "List Bullet") -> None:
        """Add a bulleted list."""
        for item in items:
            para = self.doc.add_paragraph(style=style)
            self._add_formatted_text_to_paragraph(para, item)
    
    def _add_formatted_text_to_paragraph(self, para, text: str) -> None:
        """Add formatted text to an existing paragraph, handling markdown-style formatting."""
        # Parse and format the text with bold/italic markdown
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = para.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic text  
                run = para.add_run(part[1:-1])
                run.font.italic = True
            else:
                # Normal text
                para.add_run(part)
    
    def add_numbered_list(self, items: List[str], style: str = "List Number") -> None:
        """Add a numbered list."""
        for item in items:
            para = self.doc.add_paragraph(style=style)
            self._add_formatted_text_to_paragraph(para, item)
    
    def add_code_block(self, code: str, language: str = "") -> None:
        """Add a formatted code block."""
        if language:
            lang_para = self.doc.add_paragraph(f"Code ({language}):")
            lang_para.runs[0].font.bold = True
            lang_para.runs[0].font.size = Pt(self.style.body_size - 1)
        
        # Split code into lines and add each as a separate paragraph
        code_lines = code.split('\n')
        for line in code_lines:
            code_para = self.doc.add_paragraph(line, style='Code Block')
    
    def add_table(self, data: List[List[str]], headers: Optional[List[str]] = None,
                  title: str = "") -> None:
        """Add a formatted table."""
        if title:
            self.doc.add_paragraph(title, style='Custom Heading 3')
        
        # Determine table size
        if headers:
            rows = len(data) + 1
            cols = len(headers)
        else:
            rows = len(data)
            cols = len(data[0]) if data else 0
        
        if rows == 0 or cols == 0:
            self.doc.add_paragraph("No data available for table.")
            return
        
        # Create table
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add headers if provided
        row_idx = 0
        if headers:
            header_row = table.rows[0]
            for col_idx, header in enumerate(headers):
                cell = header_row.cells[col_idx]
                cell.text = header
                # Bold header text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(self.style.body_size)
            row_idx = 1
        
        # Add data rows
        for data_row in data:
            if row_idx >= len(table.rows):
                break
            table_row = table.rows[row_idx]
            for col_idx, cell_data in enumerate(data_row):
                if col_idx < len(table_row.cells):
                    table_row.cells[col_idx].text = str(cell_data)
            row_idx += 1
    
    def add_separator(self) -> None:
        """Add a horizontal separator line."""
        para = self.doc.add_paragraph()
        para.add_run("_" * 80)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.color.rgb = RGBColor(200, 200, 200)
    
    def add_page_break(self) -> None:
        """Add a page break."""
        self.doc.add_page_break()
    
    def format_ml_research_analysis(self, analysis_data: Dict[str, Any]) -> None:
        """Format ML research analysis data into a structured document."""
        # Title section
        self.add_title(
            "ML Research Analysis Report",
            f"Query: {analysis_data.get('original_prompt', 'N/A')[:80]}...",
            add_date=True
        )
        
        # Executive Summary
        self.add_heading("Executive Summary", level=1)
        
        # Model suggestions summary
        if analysis_data.get('model_suggestions', {}).get('model_suggestions'):
            self.add_paragraph("This report provides comprehensive model recommendations based on recent research literature and domain analysis.")
            
            # Quick stats
            stats_data = [
                ["Papers Analyzed", str(analysis_data.get('model_suggestions', {}).get('papers_analyzed', 0))],
                ["Semantic Chunks", str(analysis_data.get('model_suggestions', {}).get('semantic_chunks_analyzed', 0))],
                ["Categories Detected", str(analysis_data.get('model_suggestions', {}).get('categories_considered', 0))],
                ["Validation Status", "✅ Passed" if analysis_data.get('validation_results', {}).get('validation_successful') else "❌ Failed"]
            ]
            self.add_table(stats_data, headers=["Metric", "Value"], title="Analysis Statistics")
        
        # Detected Categories
        if analysis_data.get('detected_categories'):
            self.add_heading("Detected ML Categories", level=1)
            categories = [cat.get('name', 'Unknown') for cat in analysis_data['detected_categories']]
            self.add_bullet_list(categories)
        
        # Model Recommendations
        if analysis_data.get('model_suggestions', {}).get('model_suggestions'):
            self.add_heading("Model Recommendations", level=1)
            recommendations = analysis_data['model_suggestions']['model_suggestions']
            self.add_paragraph(recommendations)
        
        # ArXiv Papers
        if analysis_data.get('arxiv_results', {}).get('papers'):
            self.add_heading("Referenced Research Papers", level=1)
            papers = analysis_data['arxiv_results']['papers'][:10]  # Top 10
            
            paper_data = []
            for i, paper in enumerate(papers, 1):
                paper_data.append([
                    str(i),
                    paper.get('title', 'Unknown')[:60] + "...",
                    str(paper.get('relevance_score', 0)),
                    paper.get('published', 'Unknown')[:10]
                ])
            
            self.add_table(paper_data, 
                          headers=["#", "Title", "Relevance", "Published"],
                          title="Top Research Papers")
        
        # Validation Results
        if analysis_data.get('validation_results'):
            self.add_heading("Validation Assessment", level=1)
            validation = analysis_data['validation_results']
            
            validation_data = [
                ["Relevance Assessment", validation.get('relevance_assessment', 'N/A')],
                ["Coverage Analysis", validation.get('coverage_analysis', 'N/A')],
                ["Quality Evaluation", validation.get('quality_evaluation', 'N/A')],
                ["Confidence Score", f"{validation.get('confidence', 0):.2f}"],
                ["Decision", validation.get('decision', 'N/A')]
            ]
            self.add_table(validation_data, headers=["Aspect", "Assessment"])
            
            if validation.get('reasoning'):
                self.add_paragraph(f"Reasoning: {validation['reasoning']}")
    
    def save(self, filename: str) -> str:
        """Save the document to a file."""
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(filename) if os.path.dirname(filename) else '.', exist_ok=True)
        
        self.doc.save(filename)
        return filename
    
    def format_ml_text_recommendations(self, text: str, title: str = "ML Model Recommendations") -> None:
        """Format ML recommendation text with proper handling of markdown-style formatting."""
        # Add title
        self.add_title(title, "Based on Recent Research Analysis")
        
        # Split text into sections
        lines = text.strip().split('\n')
        
        current_section = None
        current_model = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for introduction text
            if line.startswith('Based on') and 'recommended models' in line.lower():
                self.add_heading("Overview", level=1)
                self.add_formatted_paragraph(line)
                continue
            
            # Check for model headers (## 1., ## 2., etc.)
            model_match = re.match(r'^##\s*(\d+)\.\s*(.+)', line)
            if model_match:
                model_num = model_match.group(1)
                model_name = model_match.group(2)
                current_model = f"{model_num}. {model_name}"
                
                if current_section != 'models':
                    self.add_heading("Recommended Models", level=1)
                    current_section = 'models'
                
                self.add_heading(current_model, level=2)
                continue
            
            # Handle formatted bullet points with ** formatting
            if line.startswith('- **') and '**:' in line:
                # Parse bullet points like "- **Performance**: 42.3 mAP on COCO with 31.2 FPS"
                para = self.doc.add_paragraph(style='List Bullet')
                
                # Split by **: to separate the bold label from content
                parts = line[2:].split('**:', 1)  # Remove "- " and split on **:
                if len(parts) == 2:
                    label = parts[0].replace('**', '').strip()
                    content = parts[1].strip()
                    
                    # Add bold label
                    run = para.add_run(f"{label}: ")
                    run.font.bold = True
                    
                    # Add regular content
                    para.add_run(content)
                else:
                    # Fallback: just add the text
                    para.add_run(line[2:])
                continue
            
            # Handle regular bullet points
            if line.startswith('- '):
                para = self.doc.add_paragraph(style='List Bullet')
                self._add_formatted_text_to_paragraph(para, line[2:])
                continue
            
            # Handle lines with ** formatting but not bullet points
            if '**' in line:
                self.add_formatted_paragraph(line)
                continue
            
            # Regular paragraphs
            if line and current_section:
                self.add_paragraph(line)
        
        # Add comparison section if multiple models detected
        model_count = len(re.findall(r'^##\s*\d+\.', text, re.MULTILINE))
        if model_count > 1:
            self._add_model_comparison_section(text)
    
    def _add_model_comparison_section(self, text: str) -> None:
        """Add a comparison table section for multiple models."""
        self.add_page_break()
        self.add_heading("Model Comparison Summary", level=1)
        
        # Extract model data for comparison
        models_data = []
        model_sections = re.split(r'\n##\s*(\d+)\.\s*', text)
        
        # Process model sections
        for i in range(1, len(model_sections), 2):
            if i + 1 < len(model_sections):
                model_num = model_sections[i]
                content = model_sections[i + 1]
                
                # Extract model name
                lines = content.split('\n')
                model_name = lines[0].strip() if lines else f"Model {model_num}"
                
                # Extract performance metrics
                map_score = "N/A"
                fps_score = "N/A"
                use_case = "N/A"
                
                for line in lines:
                    if 'Performance' in line and 'mAP' in line:
                        map_match = re.search(r'(\d+\.?\d*)\s*mAP', line)
                        fps_match = re.search(r'(\d+\.?\d*)\s*FPS', line)
                        if map_match:
                            map_score = map_match.group(1)
                        if fps_match:
                            fps_score = fps_match.group(1)
                    elif 'Use Case' in line:
                        use_case = re.sub(r'.*Use Case.*?:\s*', '', line).strip()
                
                # Determine real-time capability
                try:
                    fps_val = float(fps_score) if fps_score != "N/A" else 0
                    real_time = "✅ Yes" if fps_val >= 30 else "⚠️ Marginal" if fps_val >= 20 else "❌ No"
                except:
                    real_time = "Unknown"
                
                models_data.append([
                    f"#{model_num}",
                    model_name,
                    f"{map_score}%" if map_score != "N/A" else "N/A",
                    fps_score if fps_score != "N/A" else "N/A",
                    real_time,
                    use_case[:40] + "..." if len(use_case) > 40 else use_case
                ])
        
        if models_data:
            self.add_table(
                models_data,
                headers=["Rank", "Model", "mAP", "FPS", "Real-time", "Primary Use Case"],
                title=""
            )


class MarkdownToWordConverter:
    """Convert Markdown text to formatted Word documents."""
    
    def __init__(self, formatter: Optional[WordFormatter] = None):
        """Initialize with a WordFormatter instance."""
        self.formatter = formatter or WordFormatter()
    
    def convert_markdown_text(self, markdown_text: str) -> None:
        """Convert markdown text to Word format."""
        lines = markdown_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Headers
            if line.startswith('###'):
                self.formatter.add_heading(line[3:].strip(), level=3)
            elif line.startswith('##'):
                self.formatter.add_heading(line[2:].strip(), level=2)
            elif line.startswith('#'):
                self.formatter.add_heading(line[1:].strip(), level=1)
            
            # Code blocks
            elif line.startswith('```'):
                # Find end of code block
                code_lines = []
                language = line[3:].strip()
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                self.formatter.add_code_block('\n'.join(code_lines), language)
            
            # Bullet lists
            elif line.startswith('- ') or line.startswith('* '):
                bullet_items = []
                while i < len(lines) and (lines[i].strip().startswith('- ') or 
                                        lines[i].strip().startswith('* ')):
                    bullet_items.append(lines[i].strip()[2:])
                    i += 1
                i -= 1  # Adjust for the extra increment
                
                self.formatter.add_bullet_list(bullet_items)
            
            # Regular paragraphs
            else:
                self.formatter.add_paragraph(line)
            
            i += 1


# Utility functions
def create_ml_research_report(analysis_data: Dict[str, Any], 
                             output_filename: str = None) -> str:
    """Create a formatted Word report from ML research analysis data."""
    if not output_filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"ml_research_report_{timestamp}.docx"
    
    formatter = WordFormatter()
    formatter.format_ml_research_analysis(analysis_data)
    return formatter.save(output_filename)


def convert_text_to_word(text: str, title: str = "Document", 
                        output_filename: str = None) -> str:
    """Convert plain text to a formatted Word document."""
    if not output_filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', title)[:30]
        output_filename = f"{safe_title}_{timestamp}.docx"
    
    formatter = WordFormatter()
    formatter.add_title(title)
    
    # Split text into paragraphs
    paragraphs = text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            formatter.add_paragraph(para.strip())
    
    return formatter.save(output_filename)


def convert_markdown_to_word(markdown_text: str, title: str = "Document",
                           output_filename: str = None) -> str:
    """Convert Markdown text to a formatted Word document."""
    if not output_filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', title)[:30]
        output_filename = f"{safe_title}_{timestamp}.docx"
    
    formatter = WordFormatter()
    formatter.add_title(title)
    
    converter = MarkdownToWordConverter(formatter)
    converter.convert_markdown_text(markdown_text)
    
    return formatter.save(output_filename)


# Example usage
if __name__ == "__main__":
    # Example 1: Basic document creation
    formatter = WordFormatter()
    formatter.add_title("My Research Report", "Subtitle Here")
    formatter.add_heading("Introduction", level=1)
    formatter.add_paragraph("This is an example paragraph with some content.")
    
    formatter.add_heading("Methods", level=2)
    formatter.add_bullet_list([
        "First method described here",
        "Second method with details",
        "Third approach used"
    ])
    
    formatter.add_heading("Code Example", level=2)
    formatter.add_code_block("""
def example_function():
    return "Hello, World!"

result = example_function()
print(result)
    """, "Python")
    
    filename = formatter.save("example_report.docx")
    print(f"✅ Document saved as: {filename}")
    
    # Example 2: Convert markdown
    markdown_sample = """
# Main Title

## Section 1

This is a paragraph with some content.

### Subsection

- Bullet point 1
- Bullet point 2
- Bullet point 3

## Code Section

```python
def hello():
    print("Hello, World!")
```

## Conclusion

Final thoughts here.
    """
    
    filename2 = convert_markdown_to_word(markdown_sample, "Markdown Example")
    print(f"✅ Markdown document saved as: {filename2}")
